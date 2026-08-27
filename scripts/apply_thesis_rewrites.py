#!/usr/bin/env python3
"""Apply paragraph-level rewrites to the cowpea thesis DOCX."""

from __future__ import annotations

import json
import shutil
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "r & d4 (1).docx"
OUTPUT = ROOT / "r & d4 (1)_revised.docx"
REWRITES_FILE = Path(__file__).resolve().parent / "thesis_rewrites.json"


def load_rewrites() -> dict[int, str]:
    data = json.loads(REWRITES_FILE.read_text(encoding="utf-8"))
    return {int(k): v for k, v in data.items()}


def apply_rewrites(source: Path, output: Path, rewrites: dict[int, str]) -> None:
    shutil.copy2(source, output)
    doc = Document(str(output))
    applied = 0
    for idx, new_text in sorted(rewrites.items()):
        if idx >= len(doc.paragraphs):
            print(f"Warning: paragraph index {idx} out of range; skipped")
            continue
        para = doc.paragraphs[idx]
        old = para.text.strip()
        if not old and not new_text.strip():
            continue
        para.text = new_text
        applied += 1
    doc.save(str(output))
    print(f"Applied {applied} paragraph rewrites -> {output}")


def main() -> int:
    rewrites = load_rewrites()
    apply_rewrites(SOURCE, OUTPUT, rewrites)
    return 0


if __name__ == "__main__":
    sys.exit(main())
